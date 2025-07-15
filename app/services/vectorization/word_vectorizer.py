"""
Word文档向量化器
专门处理Word文档(.doc, .docx)的内容提取和向量化，支持图片内容识别
"""

import logging
import os
import re
import tempfile
from typing import List, Dict, Any
from .base_vectorizer import BaseVectorizer
from ..enhanced_image_recognition_service import get_enhanced_image_service
from ..ocr_config_manager import get_ocr_config_manager

logger = logging.getLogger(__name__)

class WordVectorizer(BaseVectorizer):
    """Word文档向量化器，支持文本和图片内容提取"""
    
    def __init__(self):
        """初始化Word向量化器"""
        super().__init__()
        self.file_type = "word"
        # 使用统一的OCR配置管理器
        config_manager = get_ocr_config_manager()
        ocr_config = config_manager.get_resource_config_for_document_type('word')
        self.recognition_service = get_enhanced_image_service(ocr_config)
        self.config_manager = config_manager
        
        logger.info(f"✅ Word向量化器初始化完成，OCR配置: 并发{ocr_config.max_concurrent_tasks}, 超时{ocr_config.single_task_timeout}s, 最大图片{ocr_config.max_images_per_document}张")
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ['.doc', '.docx']
    
    def extract_text(self, file_path: str, extract_images: bool = True, recognition_options: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        从Word文档提取文本内容和图片内容
        
        Args:
            file_path: Word文档路径
            extract_images: 是否提取和识别图片内容
            recognition_options: 图片识别选项
            
        Returns:
            包含提取结果的字典
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return {
                    'success': False,
                    'error': f'Word文档不存在: {file_path}'
                }
            
            # 检查文件扩展名
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in self.get_supported_extensions():
                return {
                    'success': False,
                    'error': f'不支持的文件格式: {file_ext}'
                }
            
            # 设置默认识别选项
            if recognition_options is None:
                recognition_options = {
                    'engine': 'auto',
                    'languages': ['zh', 'en']
                }
            
            # 根据文件类型选择提取方法
            if file_ext == '.docx':
                return self._extract_docx(file_path, extract_images, recognition_options)
            elif file_ext == '.doc':
                return self._extract_doc(file_path, extract_images, recognition_options)
            else:
                return {
                    'success': False,
                    'error': f'不支持的Word文档格式: {file_ext}'
                }
                
        except Exception as e:
            logger.error(f"Word document text extraction failed: {e}")
            return {
                'success': False,
                'error': f'Word文档文本提取失败: {str(e)}'
            }
    
    def _extract_docx(self, file_path: str, extract_images: bool, recognition_options: Dict[str, Any]) -> Dict[str, Any]:
        """提取.docx文档内容，包括图片"""
        try:
            import docx
            from docx.document import Document
            
            doc = docx.Document(file_path)
            
            # 获取基本元数据
            metadata = {
                'file_size': os.path.getsize(file_path),
                'file_type': 'docx',
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'extraction_method': 'python-docx'
            }
            
            # 提取文档属性
            core_properties = doc.core_properties
            if core_properties:
                metadata.update({
                    'title': core_properties.title or '',
                    'author': core_properties.author or '',
                    'subject': core_properties.subject or '',
                    'keywords': core_properties.keywords or '',
                    'comments': core_properties.comments or '',
                    'created': str(core_properties.created) if core_properties.created else '',
                    'modified': str(core_properties.modified) if core_properties.modified else '',
                    'last_modified_by': core_properties.last_modified_by or '',
                    'category': core_properties.category or ''
                })
            
            # 提取内容
            content_parts = []
            image_contents = []
            total_images = 0
            
            # 基本信息
            filename = os.path.basename(file_path)
            content_parts.append(f"Word文档: {filename}")
            content_parts.append(f"文件大小: {self._format_file_size(metadata['file_size'])}")
            
            if metadata.get('title'):
                content_parts.append(f"标题: {metadata['title']}")
            if metadata.get('author'):
                content_parts.append(f"作者: {metadata['author']}")
            if metadata.get('subject'):
                content_parts.append(f"主题: {metadata['subject']}")
            if metadata.get('keywords'):
                content_parts.append(f"关键词: {metadata['keywords']}")
            
            # 提取文本内容
            text_parts = []
            
            # 提取段落文本
            for para_index, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
                
                # 处理段落中的图片（内联图片）
                if extract_images:
                    self._extract_paragraph_images(
                        paragraph, para_index, image_contents, 
                        recognition_options, total_images
                    )
            
            # 提取表格内容
            table_texts = []
            for table_index, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text:
                    table_texts.append(f"=== 表格{table_index + 1} ===\n{table_text}")
            
            # 提取文档中的图片（使用python-docx-template或直接解析）
            if extract_images:
                doc_images = self._extract_document_images(doc, recognition_options)
                image_contents.extend(doc_images)
                total_images += len(doc_images)
            
            # 构建完整内容
            if text_parts:
                content_parts.append(f"\n=== 文档正文内容 ===")
                content_parts.extend(text_parts)
            
            if table_texts:
                content_parts.append(f"\n=== 表格内容 ===")
                content_parts.extend(table_texts)
            
            if image_contents:
                content_parts.append(f"\n=== 图片识别内容 ===")
                content_parts.append(f"共识别到 {total_images} 张包含文字的图片:")
                content_parts.extend(image_contents)
            
            # 合并所有文本
            full_text = '\n'.join(text_parts + table_texts)
            comprehensive_text = '\n'.join(content_parts)
            
            # 关键词提取
            keywords = self._extract_keywords(comprehensive_text, filename, metadata)
            if keywords:
                content_parts.append(f"\n=== 检索关键词 ===")
                content_parts.append(" ".join(keywords))
                comprehensive_text = '\n'.join(content_parts)
            
            # 更新元数据
            metadata.update({
                'total_images': total_images,
                'has_images': total_images > 0,
                'has_tables': len(table_texts) > 0,
                'content_length': len(comprehensive_text),
                'text_length': len(full_text)
            })
            
            if not comprehensive_text.strip():
                return {
                    'success': False,
                    'error': 'Word文档中没有可提取的文本内容'
                }
            
            # 进行文本分块
            chunks = self.chunk_text(comprehensive_text)
            
            return {
                'success': True,
                'text': comprehensive_text,
                'metadata': metadata,
                'chunks': chunks,
                'file_type': 'word'
            }
            
        except ImportError:
            return {
                'success': False,
                'error': 'python-docx库未安装，请运行: pip install python-docx'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX文档处理失败: {str(e)}'
            }
    
    def _extract_doc(self, file_path: str, extract_images: bool, recognition_options: Dict[str, Any]) -> Dict[str, Any]:
        """提取.doc文档内容（旧版本Word）"""
        try:
            # 尝试使用python-docx2txt
            text_content = ""
            metadata = {
                'file_size': os.path.getsize(file_path),
                'file_type': 'doc',
                'extraction_method': 'unknown'
            }
            
            try:
                import docx2txt
                text_content = docx2txt.process(file_path)
                metadata['extraction_method'] = 'docx2txt'
                
            except ImportError:
                # 尝试使用antiword（需要系统安装）
                try:
                    import subprocess
                    result = subprocess.run(['antiword', file_path], 
                                          capture_output=True, text=True, encoding='utf-8')
                    
                    if result.returncode == 0 and result.stdout.strip():
                        text_content = result.stdout
                        metadata['extraction_method'] = 'antiword'
                    else:
                        # 尝试使用libreoffice转换
                        try:
                            # 使用libreoffice转换为文本
                            with tempfile.TemporaryDirectory() as temp_dir:
                                cmd = ['libreoffice', '--headless', '--convert-to', 'txt', 
                                      '--outdir', temp_dir, file_path]
                                result = subprocess.run(cmd, capture_output=True, text=True)
                                
                                if result.returncode == 0:
                                    # 查找生成的txt文件
                                    base_name = os.path.splitext(os.path.basename(file_path))[0]
                                    txt_file = os.path.join(temp_dir, f"{base_name}.txt")
                                    
                                    if os.path.exists(txt_file):
                                        with open(txt_file, 'r', encoding='utf-8') as f:
                                            text_content = f.read()
                                        metadata['extraction_method'] = 'libreoffice'
                        except Exception as e:
                            logger.warning(f"LibreOffice conversion failed: {e}")
                            
                except Exception as e:
                    logger.warning(f"External tools failed: {e}")
            
            if not text_content or not text_content.strip():
                return {
                    'success': False,
                    'error': 'DOC文档中没有可提取的文本内容，建议转换为DOCX格式'
                }
            
            # 构建内容
            filename = os.path.basename(file_path)
            content_parts = [
                f"Word文档: {filename}",
                f"文件大小: {self._format_file_size(metadata['file_size'])}",
                f"提取方法: {metadata['extraction_method']}"
            ]
            
            content_parts.append(f"\n=== 文档内容 ===")
            content_parts.append(text_content.strip())
            
            # 关键词提取
            comprehensive_text = '\n'.join(content_parts)
            keywords = self._extract_keywords(comprehensive_text, filename, metadata)
            if keywords:
                content_parts.append(f"\n=== 检索关键词 ===")
                content_parts.append(" ".join(keywords))
                comprehensive_text = '\n'.join(content_parts)
            
            # 更新元数据
            metadata.update({
                'total_images': 0,  # DOC格式难以提取图片
                'has_images': False,
                'has_tables': 'table' in text_content.lower() or '表格' in text_content,
                'content_length': len(comprehensive_text),
                'text_length': len(text_content)
            })
            
            # 进行文本分块
            chunks = self.chunk_text(comprehensive_text)
            
            return {
                'success': True,
                'text': comprehensive_text,
                'metadata': metadata,
                'chunks': chunks,
                'file_type': 'word'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'DOC文档处理失败: {str(e)}'
            }
    
    def _extract_paragraph_images(self, paragraph, para_index: int, image_contents: List[str], 
                                 recognition_options: Dict[str, Any], total_images: int) -> int:
        """从段落中提取图片内容"""
        try:
            # 检查段落中是否有内联图片
            for run in paragraph.runs:
                if hasattr(run, '_inline_graphic') or 'pic:pic' in run._element.xml:
                    # 这里可以扩展图片提取逻辑
                    # 由于python-docx对图片提取支持有限，暂时跳过
                    pass
        except Exception as e:
            logger.debug(f"Failed to extract images from paragraph {para_index}: {e}")
        
        return total_images
    
    def _extract_document_images(self, doc, recognition_options: Dict[str, Any]) -> List[str]:
        """提取文档中的图片（需要扩展库支持）"""
        image_contents = []
        
        try:
            # 这里可以使用更高级的库如python-docx2txt或直接解析docx zip文件
            # 目前保持简单实现
            pass
        except Exception as e:
            logger.debug(f"Failed to extract document images: {e}")
        
        return image_contents
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本内容"""
        try:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(" | ".join(row_data))
            
            if table_data:
                return "\n".join(table_data)
            else:
                return ""
                
        except Exception as e:
            logger.warning(f"Failed to extract table text: {e}")
            return ""
    
    def _extract_keywords(self, text: str, filename: str, metadata: Dict) -> List[str]:
        """提取关键词用于检索优化"""
        keywords = []
        
        # 文件名关键词
        base_name = os.path.splitext(filename)[0]
        filename_words = base_name.replace('_', ' ').replace('-', ' ').split()
        keywords.extend([word for word in filename_words if len(word) > 1])
        
        # 元数据关键词
        for key in ['title', 'author', 'subject', 'keywords']:
            value = metadata.get(key, '')
            if value and isinstance(value, str):
                words = value.replace('_', ' ').replace('-', ' ').split()
                keywords.extend([word for word in words if len(word) > 1])
        
        # 文本内容关键词
        if text:
            import re
            # 提取中文词汇（2个或以上连续汉字）
            chinese_words = re.findall(r'[\u4e00-\u9fff]{2,}', text)
            keywords.extend(chinese_words[:50])  # 限制数量
            
            # 提取英文单词（2个或以上字符）
            english_words = re.findall(r'[a-zA-Z]{2,}', text)
            keywords.extend(english_words[:50])  # 限制数量
            
            # 提取数字模式
            numbers = re.findall(r'\d{2,}', text)
            keywords.extend(numbers[:20])  # 限制数量
            
            # 提取可能的车牌号
            license_plates = re.findall(r'[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼使领][A-Z][A-Z0-9]{5}', text)
            keywords.extend(license_plates)
            
            # 提取可能的身份证号
            id_numbers = re.findall(r'\d{17}[\dX]', text)
            keywords.extend(id_numbers)
            
            # 提取电话号码
            phone_numbers = re.findall(r'1[3-9]\d{9}', text)
            keywords.extend(phone_numbers)
        
        # 去重并过滤
        unique_keywords = list(set(keywords))
        return [kw for kw in unique_keywords if len(kw) > 1 and not kw.isspace()][:100]  # 限制总数
    
    def _format_file_size(self, size_bytes: int) -> str:
        """格式化文件大小"""
        if size_bytes < 1024:
            return f"{size_bytes}B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f}KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f}MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.1f}GB"
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        将Word文档文本进行智能分块
        
        Args:
            text: 原始文本
            chunk_size: 块大小
            overlap: 重叠大小
            
        Returns:
            文本块列表
        """
        if not text or not text.strip():
            return []
        
        # 清理文本
        text = self._clean_text(text)
        
        chunks = []
        
        # 首先按重要段落分隔符分割
        major_sections = re.split(r'\n\s*===\s*[^=]+\s*===\s*\n', text)
        
        for section in major_sections:
            section = section.strip()
            if not section:
                continue
            
            # 如果段落本身就很长，需要进一步分割
            if len(section) <= chunk_size:
                chunks.append(section)
            else:
                # 按段落分割
                paragraphs = re.split(r'\n\s*\n', section)
                current_chunk = ""
                
                for paragraph in paragraphs:
                    paragraph = paragraph.strip()
                    if not paragraph:
                        continue
                    
                    # 如果当前段落加上当前块的长度超过chunk_size
                    if len(current_chunk) + len(paragraph) > chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                            
                            # 处理重叠
                            if overlap > 0 and len(current_chunk) > overlap:
                                overlap_text = current_chunk[-overlap:]
                                current_chunk = overlap_text + "\n" + paragraph
                            else:
                                current_chunk = paragraph
                        else:
                            # 如果单个段落就超过了chunk_size，需要强制分割
                            if len(paragraph) > chunk_size:
                                sub_chunks = self._split_long_paragraph(paragraph, chunk_size, overlap)
                                chunks.extend(sub_chunks[:-1])
                                current_chunk = sub_chunks[-1] if sub_chunks else ""
                            else:
                                current_chunk = paragraph
                    else:
                        if current_chunk:
                            current_chunk += "\n\n" + paragraph
                        else:
                            current_chunk = paragraph
                
                # 添加最后一个块
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
        
        logger.info(f"Word text split into {len(chunks)} chunks")
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """清理Word文档文本"""
        # 移除多余的空白字符
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        # 移除行末的多余空格
        text = re.sub(r' +\n', '\n', text)
        # 移除多余的空格
        text = re.sub(r' {2,}', ' ', text)
        # 移除制表符
        text = text.replace('\t', ' ')
        return text.strip()
    
    def _split_long_paragraph(self, paragraph: str, chunk_size: int, overlap: int) -> List[str]:
        """分割过长的段落"""
        chunks = []
        
        # 首先尝试按句子分割
        sentences = re.split(r'[.。!！?？;；]+', paragraph)
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            if len(current_chunk) + len(sentence) > chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    
                    # 处理重叠
                    if overlap > 0 and len(current_chunk) > overlap:
                        overlap_text = current_chunk[-overlap:]
                        current_chunk = overlap_text + sentence
                    else:
                        current_chunk = sentence
                else:
                    # 如果单个句子就超过了chunk_size，按单词分割
                    if len(sentence) > chunk_size:
                        word_chunks = self._split_by_words(sentence, chunk_size, overlap)
                        chunks.extend(word_chunks[:-1])
                        current_chunk = word_chunks[-1] if word_chunks else ""
                    else:
                        current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += sentence
                else:
                    current_chunk = sentence
        
        # 添加最后一个块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [paragraph]
    
    def _split_by_words(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """按单词分割文本"""
        words = text.split()
        chunks = []
        current_chunk = ""
        
        for word in words:
            if len(current_chunk) + len(word) + 1 > chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    
                    # 处理重叠
                    if overlap > 0:
                        chunk_words = current_chunk.split()
                        overlap_words = chunk_words[-overlap//10:] if len(chunk_words) > overlap//10 else chunk_words
                        current_chunk = " ".join(overlap_words) + " " + word
                    else:
                        current_chunk = word
                else:
                    current_chunk = word
            else:
                if current_chunk:
                    current_chunk += " " + word
                else:
                    current_chunk = word
        
        # 添加最后一个块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [text] 