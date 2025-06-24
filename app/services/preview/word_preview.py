from .base_preview import BasePreviewService
import os
import logging
import shutil
import uuid
from datetime import datetime
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import zipfile
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

class WordPreviewService(BasePreviewService):
    """Word文档预览服务"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['docx', 'doc']
        # 创建临时图片目录
        self.temp_images_dir = os.path.join('app', 'static', 'temp_images')
        os.makedirs(self.temp_images_dir, exist_ok=True)
    
    def extract_content(self, file_path, document_id=None):
        """提取Word文档内容"""
        
        # 检查是否为MCP创建的虚拟文件
        if file_path and file_path.startswith('mcp_created/'):
            return self._extract_mcp_content(file_path, document_id)
        
        # 普通文件处理
        if not self.validate_file(file_path):
            raise FileNotFoundError(f"文件不存在或无法读取: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self._extract_docx_content(file_path)
        elif file_ext == '.doc':
            return self._extract_doc_content(file_path)
        else:
            raise ValueError(f"不支持的Word文档格式: {file_ext}")
    
    def _extract_mcp_content(self, file_path, document_id):
        """提取MCP创建的Word文件内容"""
        try:
            logger.info(f"🔍 处理MCP Word文件: {file_path}")
            
            # 导入模型
            from app.models.document_models import DocumentNode, DocumentContent
            from app import db
            
            if document_id:
                # 使用document_id查找
                document = db.session.query(DocumentNode).filter_by(
                    id=document_id,
                    is_deleted=False
                ).first()
            else:
                # 使用file_path查找
                document = db.session.query(DocumentNode).filter_by(
                    file_path=file_path,
                    is_deleted=False
                ).first()
                
            if not document:
                raise FileNotFoundError(f"MCP Word文件不存在: {file_path}")
            
            # 检查虚拟文件是否存在于文件系统中
            import os
            full_virtual_path = os.path.join('uploads', file_path)
            
            if os.path.exists(full_virtual_path):
                # 文件存在于文件系统中，使用实际Word文件处理
                logger.info(f"📁 虚拟Word文件存在于文件系统: {full_virtual_path}")
                
                try:
                    file_ext = os.path.splitext(full_virtual_path)[1].lower()
                    if file_ext == '.docx':
                        return self._extract_docx_content(full_virtual_path)
                    elif file_ext == '.doc':
                        return self._extract_doc_content(full_virtual_path)
                    else:
                        raise ValueError(f"不支持的Word格式: {file_ext}")
                        
                except Exception as e:
                    logger.warning(f"处理虚拟Word文件失败，降级到文本模式: {e}")
                    # 降级到文本内容处理
                    pass
            
            # 从数据库获取文本内容（降级处理）
            logger.info(f"💾 从数据库读取MCP Word文件内容: {file_path}")
            
            content_record = db.session.query(DocumentContent).filter_by(
                document_id=document.id
            ).first()
            
            if content_record and content_record.content_text:
                # 将文本内容转换为Word预览格式
                text_content = content_record.content_text
                
                # 模拟Word文档结构
                content_data = {
                    'type': 'word',
                    'content': text_content,
                    'images': [],
                    'tables': [],
                    'metadata': {
                        'paragraphs': len([p for p in text_content.split('\n\n') if p.strip()]),
                        'words': len(text_content.split()),
                        'characters': len(text_content),
                        'source': 'mcp_created',
                        'format': 'docx',
                        'note': '📄 Word文档已生成，可通过下载功能获取完整文档文件'
                    }
                }
                
                logger.info(f"✅ MCP Word文件内容提取成功: {file_path}")
                return content_data
            else:
                # 返回空内容
                return {
                    'type': 'word',
                    'content': '[无内容]',
                    'images': [],
                    'tables': [],
                    'metadata': {
                        'paragraphs': 0,
                        'words': 0,
                        'characters': 0,
                        'source': 'mcp_created',
                        'format': 'docx',
                        'error': '无法获取文件内容'
                    }
                }
            
        except Exception as e:
            logger.error(f"❌ MCP Word文件内容提取失败: {file_path}, 错误: {str(e)}")
            raise
    
    def _extract_docx_content(self, file_path):
        """提取DOCX文档内容"""
        try:
            content_data = {
                'text': '',
                'images': [],
                'metadata': {}
            }
            
            # 使用python-docx提取内容
            doc = Document(file_path)
            
            # 提取图片
            images = self._extract_images_from_docx(file_path)
            content_data['images'] = images
            
            # 提取文本内容，包括表格
            content_parts = []
            
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # 段落
                    paragraph = Paragraph(element, doc)
                    text = paragraph.text.strip()
                    if text:
                        content_parts.append(text)
                elif isinstance(element, CT_Tbl):
                    # 表格
                    table = Table(element, doc)
                    table_text = self._extract_table_text(table)
                    if table_text:
                        content_parts.append(f"\n[表格内容]\n{table_text}\n[表格结束]\n")
            
            content_data['text'] = '\n\n'.join(content_parts)
            
            # 提取元数据
            core_props = doc.core_properties
            content_data['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0
            }
            
            # 统计信息
            content_data['metadata'].update({
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables),
                'sections_count': len(doc.sections),
                'images_count': len(images)
            })
            
            logger.info(f"✅ DOCX内容提取成功: {file_path}, 图片数量: {len(images)}")
            return content_data
            
        except Exception as e:
            logger.error(f"❌ DOCX内容提取失败: {file_path}, 错误: {str(e)}")
            raise
    
    def _extract_table_text(self, table):
        """提取表格文本内容"""
        table_rows = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text if cell_text else '空')
            table_rows.append(' | '.join(row_cells))
        return '\n'.join(table_rows)
    
    def _extract_images_from_docx(self, file_path):
        """从DOCX文档中提取图片"""
        images = []
        try:
            # 生成唯一的图片前缀
            image_prefix = str(uuid.uuid4())[:8]
            
            # 解压docx文件
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # 查找媒体文件夹中的图片
                media_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                
                for i, media_file in enumerate(media_files):
                    # 检查是否是图片文件
                    if any(media_file.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']):
                        # 提取文件扩展名
                        file_ext = os.path.splitext(media_file)[1]
                        
                        # 生成唯一的文件名
                        temp_filename = f"{image_prefix}_img_{i+1}{file_ext}"
                        temp_path = os.path.join(self.temp_images_dir, temp_filename)
                        
                        # 提取图片文件
                        with zip_file.open(media_file) as source, open(temp_path, 'wb') as target:
                            target.write(source.read())
                        
                        # 生成访问URL
                        image_url = f"/static/temp_images/{temp_filename}"
                        
                        images.append({
                            'url': image_url,
                            'filename': temp_filename,
                            'original_name': os.path.basename(media_file),
                            'index': i + 1
                        })
                        
                        logger.info(f"提取图片: {media_file} -> {temp_filename}")
            
        except Exception as e:
            logger.error(f"图片提取失败: {file_path}, 错误: {str(e)}")
        
        return images
    
    def _extract_doc_content(self, file_path):
        """提取DOC文档内容（老格式）"""
        try:
            # 对于.doc格式，尝试使用python-docx2txt
            try:
                import docx2txt
                
                # 注意：docx2txt只能处理基于ZIP的文档格式，不能处理老的.doc格式
                # 如果文件是老的.doc格式，会抛出zipfile.BadZipFile异常
                text_content = docx2txt.process(file_path)
                
                # 尝试提取图片
                images = []
                try:
                    # docx2txt可以提取图片到指定目录
                    image_prefix = str(uuid.uuid4())[:8]
                    temp_dir = os.path.join(self.temp_images_dir, image_prefix)
                    os.makedirs(temp_dir, exist_ok=True)
                    
                    # 提取图片
                    docx2txt.process(file_path, temp_dir)
                    
                    # 查找提取的图片
                    for i, filename in enumerate(os.listdir(temp_dir)):
                        if any(filename.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']):
                            # 移动图片到主目录
                            old_path = os.path.join(temp_dir, filename)
                            new_filename = f"{image_prefix}_img_{i+1}_{filename}"
                            new_path = os.path.join(self.temp_images_dir, new_filename)
                            shutil.move(old_path, new_path)
                            
                            images.append({
                                'url': f"/static/temp_images/{new_filename}",
                                'filename': new_filename,
                                'original_name': filename,
                                'index': i + 1
                            })
                    
                    # 清理临时目录
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
                        
                except Exception as img_error:
                    logger.warning(f"DOC图片提取失败: {str(img_error)}")
                
                content_data = {
                    'text': text_content or '无法提取文本内容',
                    'images': images,
                    'metadata': {
                        'note': 'DOC格式支持，建议转换为DOCX格式以获取更好体验',
                        'images_count': len(images)
                    }
                }
                
                logger.info(f"✅ DOC内容提取成功: {file_path}, 图片数量: {len(images)}")
                return content_data
                
            except ImportError:
                # 如果没有docx2txt库，返回基本信息
                content_data = {
                    'text': '需要安装docx2txt库以支持.doc格式的内容提取\n请运行: pip install docx2txt',
                    'images': [],
                    'metadata': {
                        'note': '需要安装docx2txt库以获取完整支持',
                        'images_count': 0
                    }
                }
                
                logger.warning(f"⚠️ DOC格式需要docx2txt库: {file_path}")
                return content_data
            except Exception as docx2txt_error:
                # docx2txt无法处理这个文件（可能是老的.doc格式）
                error_msg = str(docx2txt_error)
                if "not a zip file" in error_msg.lower() or "badzipfile" in error_msg.lower():
                    # 这是老的.doc格式，docx2txt无法处理
                    content_data = {
                        'text': '检测到老式.doc格式文档。\n\n由于技术限制，暂时无法直接预览老式.doc格式的内容。\n\n建议解决方案：\n1. 使用Microsoft Word将文档另存为.docx格式\n2. 使用在线转换工具将.doc转换为.docx格式\n3. 或直接下载原文件在本地软件中查看',
                        'images': [],
                        'metadata': {
                            'note': '老式.doc格式，建议转换为.docx格式以获取预览功能',
                            'format': 'Legacy DOC format',
                            'images_count': 0,
                            'suggestion': '请将文档转换为.docx格式以获得更好的预览体验'
                        }
                    }
                    
                    logger.info(f"⚠️ 检测到老式DOC格式: {file_path}")
                    return content_data
                else:
                    # 其他错误，重新抛出
                    raise docx2txt_error
                
        except Exception as e:
            logger.error(f"❌ DOC内容提取失败: {file_path}, 错误: {str(e)}")
            raise
    
    def get_metadata(self, file_path):
        """获取Word文档元数据"""
        if not self.validate_file(file_path):
            raise FileNotFoundError(f"文件不存在或无法读取: {file_path}")
        
        try:
            metadata = {
                'file_size': self.get_file_size(file_path),
                'file_size_formatted': self.format_file_size(self.get_file_size(file_path)),
                'file_type': 'Word Document',
                'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
            }
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.docx':
                try:
                    doc = Document(file_path)
                    core_props = doc.core_properties
                    
                    metadata.update({
                        'title': core_props.title or '',
                        'author': core_props.author or '',
                        'subject': core_props.subject or '',
                        'keywords': core_props.keywords or '',
                        'comments': core_props.comments or '',
                        'created': core_props.created.isoformat() if core_props.created else '',
                        'modified': core_props.modified.isoformat() if core_props.modified else '',
                        'last_modified_by': core_props.last_modified_by or '',
                        'revision': core_props.revision or 0,
                        'paragraphs_count': len(doc.paragraphs),
                        'tables_count': len(doc.tables),
                        'sections_count': len(doc.sections)
                    })
                except Exception as e:
                    logger.warning(f"无法提取DOCX详细元数据: {file_path}, 错误: {str(e)}")
                    metadata['error'] = str(e)
            else:
                metadata['format_note'] = 'DOC格式，建议转换为DOCX以获取更多信息'
            
            return metadata
            
        except Exception as e:
            logger.error(f"获取Word元数据失败: {file_path}, 错误: {str(e)}")
            return {
                'file_size': self.get_file_size(file_path),
                'file_size_formatted': self.format_file_size(self.get_file_size(file_path)),
                'file_type': 'Word Document',
                'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
                'error': str(e)
            }
    
    def generate_thumbnail(self, file_path, output_path=None, size=(200, 200)):
        """生成Word文档缩略图"""
        # Word文档缩略图生成需要特殊工具
        # 目前返回None，表示不支持
        logger.info(f"Word文档缩略图生成暂不支持: {file_path}")
        return None
    
    def cleanup_temp_images(self, image_prefix=None):
        """清理临时图片文件"""
        try:
            if image_prefix:
                # 清理特定前缀的图片
                for filename in os.listdir(self.temp_images_dir):
                    if filename.startswith(image_prefix):
                        file_path = os.path.join(self.temp_images_dir, filename)
                        os.remove(file_path)
                        logger.info(f"清理临时图片: {filename}")
            else:
                # 清理所有超过1小时的临时图片
                import time
                current_time = time.time()
                for filename in os.listdir(self.temp_images_dir):
                    file_path = os.path.join(self.temp_images_dir, filename)
                    if current_time - os.path.getmtime(file_path) > 3600:  # 1小时
                        os.remove(file_path)
                        logger.info(f"清理过期图片: {filename}")
        except Exception as e:
            logger.error(f"清理临时图片失败: {str(e)}") 